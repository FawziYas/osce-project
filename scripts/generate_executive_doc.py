"""
Generate a professional executive-level Word document (.docx)
describing the OSCE Exam Platform for academic leadership review.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1A, 0x1A, 0x2E)   # #1A1A2E
TEAL       = RGBColor(0x00, 0x7B, 0x83)   # #007B83
GOLD       = RGBColor(0xC8, 0x9E, 0x3F)   # #C89E3F
LIGHT_GREY = RGBColor(0xF2, 0xF4, 0xF7)   # #F2F4F7
MID_GREY   = RGBColor(0x60, 0x60, 0x60)   # #606060
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)

REPORT_DATE = datetime.date.today().strftime("%B %d, %Y")
VERSION     = "v3.2"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc, color_hex: str = '007B83', thickness_pt: int = 12):
    """Add a thin coloured border below the last paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'),  str(thickness_pt))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading(doc, text: str, level: int = 1):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = NAVY if level == 1 else TEAL
    if level == 1:
        add_horizontal_rule(doc)
    return p

def body(doc, text: str, indent: bool = False):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    run = p.runs[0]
    run.font.size      = Pt(10.5)
    run.font.color.rgb = MID_GREY
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    return p

def bullet(doc, items: list, bold_prefix: str = None):
    """Add a bullet list."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(1)
        if bold_prefix and item.startswith(bold_prefix):
            pass
        # Support "Bold Part: rest" splitting
        if ':' in item:
            parts = item.split(':', 1)
            r1 = p.add_run(parts[0] + ':')
            r1.bold           = True
            r1.font.size      = Pt(10.5)
            r1.font.color.rgb = NAVY
            r2 = p.add_run(parts[1])
            r2.font.size      = Pt(10.5)
            r2.font.color.rgb = MID_GREY
        else:
            r = p.add_run(item)
            r.font.size      = Pt(10.5)
            r.font.color.rgb = MID_GREY

def two_column_table(doc, rows: list, header: tuple = None):
    """
    Render a clean two-column table.
    rows = list of (left, right) tuples.
    """
    col_count = 2
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=col_count)
    table.style           = 'Table Grid'
    table.alignment       = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit   = True

    row_offset = 0
    if header:
        hrow = table.rows[0]
        for ci, hdr in enumerate(header):
            cell = hrow.cells[ci]
            set_cell_bg(cell, '1A1A2E')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.bold           = True
            run.font.size      = Pt(10)
            run.font.color.rgb = WHITE
            p.alignment        = WD_ALIGN_PARAGRAPH.CENTER
        row_offset = 1

    for ri, (left, right) in enumerate(rows):
        row = table.rows[ri + row_offset]
        bg  = 'F2F4F7' if ri % 2 == 0 else 'FFFFFF'
        lc  = row.cells[0]
        rc  = row.cells[1]

        set_cell_bg(lc, bg)
        set_cell_bg(rc, bg)

        lp = lc.paragraphs[0]
        l_run = lp.add_run(left)
        l_run.bold           = True
        l_run.font.size      = Pt(10)
        l_run.font.color.rgb = NAVY

        rp = rc.paragraphs[0]
        r_run = rp.add_run(right)
        r_run.font.size      = Pt(10)
        r_run.font.color.rgb = MID_GREY

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(6)
        row.cells[1].width = Cm(10)

    doc.add_paragraph()  # spacer

def add_cover_page(doc):
    """Build a visually dominant cover page."""
    # Top accent bar via a 1-row, 1-col table
    bar = doc.add_table(rows=1, cols=1)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = bar.rows[0].cells[0]
    set_cell_bg(cell, '1A1A2E')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('  EXECUTIVE BRIEFING DOCUMENT  ')
    r.bold           = True
    r.font.size      = Pt(11)
    r.font.color.rgb = GOLD

    doc.add_paragraph()

    # Institution label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('COLLEGE OF MEDICINE')
    r.bold           = True
    r.font.size      = Pt(12)
    r.font.color.rgb = TEAL

    doc.add_paragraph()
    doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('OSCE Examination Management\nPlatform')
    r.bold           = True
    r.font.size      = Pt(28)
    r.font.color.rgb = NAVY

    doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Enterprise-Grade Digital Infrastructure for Objective Structured\n'
        'Clinical Examination Administration in Academic Medicine'
    )
    r.italic         = True
    r.font.size      = Pt(13)
    r.font.color.rgb = TEAL

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Meta table
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ('Prepared for',  'Dean and Academic Leadership Board'),
        ('Classification','CONFIDENTIAL — Internal Use Only'),
        ('Document Date', REPORT_DATE),
        ('Version',       VERSION),
    ]
    for ri, (k, v) in enumerate(meta_data):
        bg = 'F2F4F7' if ri % 2 == 0 else 'FFFFFF'
        lc = meta.rows[ri].cells[0]
        rc = meta.rows[ri].cells[1]
        set_cell_bg(lc, bg)
        set_cell_bg(rc, bg)
        lp = lc.paragraphs[0]
        lr = lp.add_run(k)
        lr.bold           = True
        lr.font.size      = Pt(10)
        lr.font.color.rgb = NAVY
        rp = rc.paragraphs[0]
        rr = rp.add_run(v)
        rr.font.size      = Pt(10)
        rr.font.color.rgb = MID_GREY

    for row in meta.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    doc.add_page_break()


# ── Main document builder ─────────────────────────────────────────────────────

def build_document(output_path: str):
    doc = Document()

    # ── Page setup (A4, 2.5 cm margins) ──────────────────────────────────────
    section = doc.sections[0]
    section.page_height   = Cm(29.7)
    section.page_width    = Cm(21.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name          = 'Calibri'
    style.font.size          = Pt(10.5)
    style.font.color.rgb     = MID_GREY

    # ── Cover page ────────────────────────────────────────────────────────────
    add_cover_page(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # 1. EXECUTIVE SUMMARY
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '1.  EXECUTIVE SUMMARY')
    body(doc,
        'This document presents a comprehensive technical and strategic overview of the OSCE '
        'Examination Management Platform — a purpose-built, enterprise-grade digital solution '
        'developed exclusively for the College of Medicine to modernise, secure, and scale '
        'the administration of Objective Structured Clinical Examinations (OSCE).'
    )
    body(doc,
        'The platform addresses the full operational lifecycle of OSCE delivery: from course '
        'and station authoring through multi-path session scheduling, real-time examiner '
        'scoring (with offline resilience), to automated report generation and institutional '
        'analytics. It is engineered on a hardened, production-grade technology stack with '
        'documented security controls, performance optimisations, and a scalable modular '
        'architecture capable of supporting departmental growth without architectural revision.'
    )
    body(doc,
        'The system is presently deployed and operational, supporting live examination '
        'sessions across two student cohorts (Year 4 and Year 6), with a documented audit '
        'trail for all critical interactions and compliance-ready data governance controls.'
    )

    # ─────────────────────────────────────────────────────────────────────────
    # 2. SYSTEM OVERVIEW AND OBJECTIVES
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '2.  SYSTEM OVERVIEW AND OBJECTIVES')
    body(doc,
        'The OSCE Examination Management Platform is a full-stack web application that '
        'centralises every dimension of OSCE administration under a single, role-governed '
        'interface. The system eliminates paper-based marking, unstructured spreadsheet '
        'workflows, and manual score aggregation — replacing them with real-time digital '
        'processes that enforce institutional assessment policy by design.'
    )

    heading(doc, 'Strategic Objectives', level=2)
    bullet(doc, [
        'Operational Efficiency: Eliminate administrative overhead associated with manual OSCE coordination, examiner briefings, and post-session score consolidation.',
        'Assessment Integrity: Enforce consistent structured marking criteria across all paths, examiners, and cohorts through system-level controls.',
        'Data Sovereignty: Retain all examination data within institution-controlled infrastructure, with full audit accountability.',
        'Regulatory Readiness: Maintain structured records aligned with accreditation body requirements and academic quality assurance standards.',
        'Scalability: Support expanded session volumes, additional campuses, and increased student cohort sizes without platform re-engineering.',
    ])

    # ─────────────────────────────────────────────────────────────────────────
    # 3. CORE FEATURES AND FUNCTIONAL CAPABILITIES
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '3.  CORE FEATURES AND FUNCTIONAL CAPABILITIES')

    heading(doc, '3.1  Course and Examination Authoring', level=2)
    bullet(doc, [
        'Course Management: Structured course records linked to Intended Learning Outcomes (ILOs), theme categories, and OSCE mark weightings.',
        'Station Library: Centralised repository of reusable station templates with binary, partial-credit, and scale-based checklist items.',
        'ILO Linkage: Every checklist item is mapped to a specific ILO, enabling automated ILO-level performance analytics.',
        'Criticality Flags: Individual checklist items may be designated as critical, allowing fail-override logic at the institutional level.',
    ])

    heading(doc, '3.2  Session Planning and Multi-Path Rotation', level=2)
    bullet(doc, [
        'Multi-Path Circuits: Each examination session supports configurable circular rotation paths (e.g., 5–10 paths per session), replicating identical stations across all paths to ensure marking parity.',
        'Student Assignment: Students are assigned to specific paths with sequenced station progression, automatically redistributed upon roster changes.',
        'Examiner Assignment: Examiners are paired to specific stations per path per session, with system-level enforcement preventing cross-assignment scoring.',
        'Session Status Lifecycle: Sessions progress through defined states (Scheduled → In Progress → Completed → Locked) with role-governed transition controls.',
    ])

    heading(doc, '3.3  Real-Time and Offline Examiner Scoring', level=2)
    bullet(doc, [
        'Dedicated Examiner Interface: Examiners access a purpose-built, mobile-responsive scoring interface optimised for tablet and handheld use during live sessions.',
        'Progressive Web Application (PWA): The examiner module is engineered as an offline-capable PWA, enabling full scoring functionality in low-connectivity examination venues with automatic background synchronisation upon reconnection.',
        'Assignment Enforcement: The scoring engine validates examiner–station–session assignment on every submission, rejecting unauthorised score entries at the API layer.',
        'Correction Workflow: Designated supervisors may unlock individual station scores for authorised correction, with full audit recording of pre- and post-correction values.',
    ])

    heading(doc, '3.4  Reporting and Analytics', level=2)
    bullet(doc, [
        'Session Result Reports: Per-session reports displaying individual student totals, station-level scores, pass/fail status, and percentage against maximum.',
        'ILO Score Export (XLSX): Machine-readable Excel exports mapping each student\'s performance against individual ILOs across all session stations.',
        'PDF Session Reports: Print-ready A4 PDF documents with institution header, examiner assignment tables, path and station breakdowns, and full student score listings — with Arabic text rendering support.',
        'CSV Raw Data Export: Structured raw score export for integration with institutional MIS or statistical analysis tools.',
        'Soft-Delete Audit Trail: Deleted entities retain historical associations with session scores, preserving reporting integrity post-deletion.',
    ])

    # ─────────────────────────────────────────────────────────────────────────
    # 4. USER EXPERIENCE AND INTERFACE DESIGN
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '4.  USER EXPERIENCE AND INTERFACE DESIGN')
    body(doc,
        'The platform\'s interface is engineered across three distinct operational contexts — '
        'administrative coordination, examination delivery, and executive reporting — each '
        'presenting a role-appropriate experience without surfacing irrelevant controls or '
        'sensitive system components.'
    )
    bullet(doc, [
        'Responsive Design: Full mobile and tablet responsiveness ensures usability across device classes used during live examination delivery.',
        'Role-Isolated Dashboards: Coordinators, administrators, and examiners each receive contextually appropriate views; no role exposes another\'s operational controls.',
        'Bilingual Support: Full Arabic text rendering via native browser HarfBuzz/DirectWrite pipelines with RTL layout preservation in all exported documents.',
        'Session-Aware UI: Interface elements dynamically reflect session state — scoring controls are visually and functionally disabled outside active session windows.',
        'Force-Change-Password Flow: First-login users are redirected to a mandatory password change screen before accessing any system functionality.',
        'Activity-Based Timeouts: Administrative roles (5-minute inactivity timeout) and examiner roles (30-minute inactivity timeout) enforce session expiry appropriate to each operational context.',
    ])

    # ─────────────────────────────────────────────────────────────────────────
    # 5. SECURITY ARCHITECTURE
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '5.  SECURITY ARCHITECTURE')
    body(doc,
        'The platform has undergone a formal internal security audit (documented in '
        'SECURITY_AUDIT_REPORT.md). The following controls are implemented and verified '
        'as part of the production deployment baseline.'
    )

    heading(doc, '5.1  Authentication and Access Control', level=2)
    bullet(doc, [
        'Four-Tier Role Hierarchy: Superuser → Admin → Coordinator → Examiner, with each tier governing a strictly scoped set of system capabilities.',
        'Rate-Limited Login: Brute-force protection via django-axes, enforcing configurable lockout thresholds after failed authentication attempts.',
        'Admin Token Scoping: The secret administrative gateway URL is injected exclusively into authenticated staff-level sessions, never exposed in examiner-facing templates.',
        'Django Admin Restriction: Django\'s administrative interface is accessible only to superusers; coordinator and examiner accounts are explicitly excluded.',
    ])

    heading(doc, '5.2  API and Data Layer Protection', level=2)
    bullet(doc, [
        'Middleware-Level Path Enforcement: Middleware blocks all /api/creator/ and /creator/ endpoints for non-coordinator roles at the request layer, prior to view execution.',
        'Examiner Assignment Verification: Every scoring API endpoint validates that the requesting examiner is explicitly assigned to the target station–session combination before processing any data.',
        'Session Status Validation: Score submission endpoints reject requests when the target session is in any state other than in_progress, preventing retroactive or premature scoring.',
        'CSRF Protection: All form submissions and AJAX scoring calls enforce Django\'s CSRF token framework; previously exempt endpoints have been hardened.',
        'JSON Body Validation: All 13+ JSON request parsing operations are wrapped in structured exception handlers, returning controlled 400 Bad Request responses on malformed input.',
        'Sync Ownership Verification: The offline data synchronisation endpoint validates that the submitting examiner owns every existing score record before permitting update operations.',
    ])

    heading(doc, '5.3  Audit and Compliance', level=2)
    bullet(doc, [
        'Comprehensive Audit Logging: All critical operations — user authentication, score recording, administrative data modifications — are persisted to a structured audit log with actor, timestamp, and affected entity.',
        'Login Audit Trail: Login events, failed attempts, and session terminations are captured with IP address and user agent metadata.',
        'Soft Deletion with Preservation: Deleted entities (examiners, stations) retain referential integrity with all associated score records, ensuring historical report accuracy.',
        'Score Immutability by Default: Submitted scores are locked at the session level; correction workflows require explicit supervisor authorisation and generate a corresponding audit record.',
    ])

    # ─────────────────────────────────────────────────────────────────────────
    # 6. SYSTEM PERFORMANCE AND SCALABILITY
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '6.  SYSTEM PERFORMANCE AND SCALABILITY')
    body(doc,
        'The platform has been subjected to a formal performance audit (documented in '
        'PERFORMANCE_OPTIMIZATION_REPORT.md) targeting query efficiency, memory allocation, '
        'and concurrency under realistic examination-day load profiles.'
    )

    heading(doc, '6.1  Query and Database Optimisations', level=2)
    bullet(doc, [
        'N+1 Elimination: All major API endpoints have been refactored to eliminate N+1 query patterns through select_related(), prefetch_related(), and annotate() strategies, reducing per-request query counts from O(n) to O(1).',
        'Database Indexes: Status and foreign-key columns on high-frequency query paths (ExamSession.status, SessionStudent.status) carry explicit database indexes.',
        'Bulk Operations: Multi-record operations (e.g., student redistribution) use bulk_update() rather than per-row save() calls, reducing transaction overhead by an order of magnitude.',
        'Paginated API Responses: All listing endpoints implement server-side pagination, preventing unbounded result sets from generating excessive memory allocation under large cohort sizes.',
    ])

    heading(doc, '6.2  Scalability Architecture', level=2)
    bullet(doc, [
        'Database-Agnostic ORM: The data layer is abstracted through Django\'s ORM, enabling migration from SQLite (development) to PostgreSQL (production) or enterprise database clusters without application-layer changes.',
        'Redis Caching Layer: Production deployment incorporates a Redis caching backend (django-redis) for session state and high-frequency read endpoints.',
        'Stateless Application Layer: The application server is stateless, enabling horizontal scaling across multiple Gunicorn worker processes or container replicas.',
        'WhiteNoise Static File Serving: Static assets are served via WhiteNoise with long-term cache headers and gzip compression, eliminating static-file load from the application process.',
        'ASGI-Ready: The platform exposes an ASGI interface, positioning it for future WebSocket-based real-time features without architectural migration.',
    ])

    # ─────────────────────────────────────────────────────────────────────────
    # 7. DATA MANAGEMENT AND BACKUP STRATEGY
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '7.  DATA MANAGEMENT AND BACKUP STRATEGY')
    bullet(doc, [
        'Structured Relational Data Model: All examination data — courses, ILOs, stations, sessions, scores, and audit records — is persisted in a normalised relational schema with referential integrity constraints.',
        'PostgreSQL Production Backend: The production environment targets PostgreSQL with binary driver (psycopg2), supporting point-in-time recovery, WAL archiving, and streaming replication when configured at the infrastructure level.',
        'Soft Deletion Architecture: No examination-critical entity is hard-deleted without explicit superuser confirmation; soft-deleted records remain queryable for historical reporting.',
        'Environment-Isolated Configuration: All credentials, database connection strings, secret keys, and environment-specific flags are managed through environment variables (django-environ), ensuring no sensitive configuration is embedded in the codebase.',
        'Sentry Error Monitoring: Production deployments integrate Sentry (sentry-sdk[django]) for real-time exception tracking, enabling proactive incident detection and resolution before data integrity is affected.',
        'Export Redundancy: Critical session data can be exported to XLSX and CSV formats at any time, providing a human-readable data extract independent of the database engine.',
    ])

    # ─────────────────────────────────────────────────────────────────────────
    # 8. TECHNOLOGY STACK
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '8.  TECHNOLOGY STACK')
    two_column_table(
        doc,
        header=('Component', 'Technology / Library'),
        rows=[
            ('Application Framework',      'Django 5.2.11 (Python 3.13)'),
            ('Production Database',        'PostgreSQL 15+ via psycopg2-binary 2.9.11'),
            ('Development Database',       'SQLite 3 (schema-identical to production)'),
            ('Caching Layer',              'Redis 7 via django-redis 5.4.0'),
            ('WSGI / ASGI Server',         'Gunicorn 25.1.0 (ASGI via Uvicorn optional)'),
            ('Static File Serving',        'WhiteNoise 6.11.0 with gzip compression'),
            ('Authentication Protection',  'django-axes 8.3.1 (brute-force lockout)'),
            ('Error Monitoring',           'Sentry SDK 2.19.2 (sentry-sdk[django])'),
            ('Excel Export',               'openpyxl 3.1.5'),
            ('PDF Generation',             'ReportLab 4.4.10 + Browser-Native Print'),
            ('Arabic / RTL Rendering',     'arabic-reshaper 3.0.0 + python-bidi 0.6.7'),
            ('Environment Configuration',  'django-environ 0.13.0'),
            ('Offline PWA Support',        'Service Worker API + Background Sync'),
            ('Frontend Framework',         'Bootstrap 5.3 + Bootstrap Icons 1.11'),
            ('Deployment Target',          'Linux / Azure App Service / Docker'),
        ]
    )

    # ─────────────────────────────────────────────────────────────────────────
    # 9. FUTURE EXPANSION AND SUSTAINABILITY
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '9.  FUTURE EXPANSION AND SUSTAINABILITY')
    body(doc,
        'The modular architecture of the platform is designed to accommodate institutional '
        'growth and feature evolution without requiring foundational re-engineering. '
        'The following capability expansions are architecturally supported within the '
        'existing system boundaries:'
    )
    bullet(doc, [
        'Multi-Departmental Multi-Tenancy: Department-level data partitioning to support parallel OSCE programmes across multiple clinical specialties or campuses.',
        'Automated Student Import: Batch student enrolment via integration with institutional Student Information Systems (SIS) through structured CSV/API interfaces.',
        'Real-Time Supervision Dashboard: Live session monitoring panels for supervisors displaying per-path scoring progress, anomaly flags, and examiner connectivity status via WebSocket channels (ASGI infrastructure already in place).',
        'LMS Integration: Standards-based Learning Management System integration (LTI / REST) for grade passback and student portfolio linkage.',
        'Advanced Analytics Module: Cohort-level ILO performance trending, inter-rater reliability scoring, and longitudinal assessment analytics for programme-level quality review.',
        'CBME Alignment Reporting: Competency-Based Medical Education outcome mapping reports to support national licensing board and accreditation submissions.',
        'Mobile Native Applications: Native iOS and Android examiner applications leveraging the existing REST API layer, with enhanced biometric authentication.',
    ])

    # ─────────────────────────────────────────────────────────────────────────
    # 10. STRATEGIC VALUE TO THE INSTITUTION
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, '10.  STRATEGIC VALUE TO THE INSTITUTION')
    body(doc,
        'The OSCE Examination Management Platform represents a strategic investment in '
        'institutional digital infrastructure that delivers measurable value across '
        'operational, academic, and compliance dimensions.'
    )

    two_column_table(
        doc,
        header=('Strategic Dimension', 'Institutional Benefit'),
        rows=[
            ('Operational Excellence',
             'Eliminates manual coordination, paper-based marking, and post-session '
             'data consolidation — reducing administrative labour per session significantly.'),
            ('Assessment Integrity',
             'Enforces standardised marking criteria system-wide, eliminating inter-examiner '
             'inconsistency arising from unstructured marking forms. All deviations are auditable.'),
            ('Accreditation Readiness',
             'Structured ILO linkage, audit trails, and exportable performance data provide '
             'documentation artefacts directly applicable to national and international '
             'accreditation submissions (e.g., NCAAA, WFME, LCME).'),
            ('Data-Driven Quality Assurance',
             'ILO-level performance analytics enable curriculum committees to identify '
             'competency gaps at cohort level, driving evidence-based programme improvement.'),
            ('Risk Mitigation',
             'System-enforced access controls, audit logging, and soft-deletion policies '
             'mitigate legal and regulatory risk associated with score disputes and data loss.'),
            ('Institutional Ownership',
             'The platform is developed and maintained as internal institutional property — '
             'eliminating ongoing vendor licence costs and dependency on third-party exam '
             'delivery platforms whose data policies may not align with institutional governance.'),
            ('Student Experience',
             'Faster result dissemination, structured feedback potential, and bilingual '
             'interface support contribute to an equitable and professionally delivered '
             'assessment experience.'),
        ]
    )

    # ─────────────────────────────────────────────────────────────────────────
    # Closing statement
    # ─────────────────────────────────────────────────────────────────────────
    heading(doc, 'CLOSING STATEMENT')
    body(doc,
        'The OSCE Examination Management Platform constitutes a mature, production-ready '
        'digital infrastructure asset. It has been designed, audited, and deployed with '
        'institutional-grade standards of security, performance, and extensibility. '
        'Its modular architecture ensures that the platform will continue to serve the '
        'institution\'s evolving assessment requirements without incurring the cost and '
        'disruption of periodic system replacements.'
    )
    body(doc,
        'The academic leadership board is invited to review the supplementary technical '
        'documentation — including the Security Audit Report and Performance Optimisation '
        'Report — for a detailed record of the engineering controls and optimisation '
        'measures applied to this platform.'
    )

    add_horizontal_rule(doc, color_hex='C89E3F', thickness_pt=18)

    # Footer note
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(
        f'CONFIDENTIAL  —  Prepared {REPORT_DATE}  —  {VERSION}  —'
        '  OSCE Examination Management Platform  '
        '  College of Medicine'
    )
    r.italic         = True
    r.font.size      = Pt(8.5)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f'[✓] Document saved: {output_path}')


if __name__ == '__main__':
    base_dir    = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    output_path = os.path.join(base_dir, 'OSCE_Platform_Executive_Briefing.docx')
    build_document(output_path)
