"""
Generate: OSCE_Workflow_Guide.docx
Focused document covering Examiner Workflow and Creator/Coordinator Workflow.
Run: python scripts/generate_workflow_doc.py
"""
import os, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x1A, 0x2E)
TEAL   = RGBColor(0x00, 0x85, 0x80)
AMBER  = RGBColor(0xC0, 0x7D, 0x00)
GREY   = RGBColor(0x44, 0x44, 0x55)
LGREY  = RGBColor(0x88, 0x88, 0x99)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Low-level helpers ─────────────────────────────────────────────────────────

def _bg(cell, hex6):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)


def _rule(doc, hex6='008580', sz=6):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b    = OxmlElement('w:bottom')
    b.set(qn('w:val'),   'single')
    b.set(qn('w:sz'),    str(sz))
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), hex6)
    pBdr.append(b)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


# ── Semantic helpers ──────────────────────────────────────────────────────────

def chapter(doc, text, teal=True):
    """Top-level chapter heading."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(15)
    run.font.color.rgb = TEAL if teal else NAVY
    run.font.name      = 'Calibri'
    _rule(doc, '008580' if teal else '1A1A2E', sz=8)
    return p


def section(doc, text):
    """Second-level heading."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(12)
    run.font.color.rgb = NAVY
    run.font.name      = 'Calibri'
    return p


def subsection(doc, text):
    """Third-level heading."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(10.5)
    run.font.color.rgb = TEAL
    run.font.name      = 'Calibri'
    return p


def para(doc, text, colour=None):
    """Body paragraph."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = colour or GREY
    run.font.name      = 'Calibri'
    return p


def bul(doc, items):
    """
    Bullet list. Each item may be:
      'plain text'          → plain bullet
      ('Bold label', 'rest')→ bold prefix + normal text
    """
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        if isinstance(item, tuple):
            label, rest = item
            r1 = p.add_run(label + '  ')
            r1.bold           = True
            r1.font.size      = Pt(10.5)
            r1.font.color.rgb = NAVY
            r1.font.name      = 'Calibri'
            r2 = p.add_run(rest)
            r2.font.size      = Pt(10.5)
            r2.font.color.rgb = GREY
            r2.font.name      = 'Calibri'
        else:
            r = p.add_run(item)
            r.font.size      = Pt(10.5)
            r.font.color.rgb = GREY
            r.font.name      = 'Calibri'


def workflow_table(doc, steps):
    """
    Numbered workflow step table.
    steps = list of (step_num_str, title, description)
    """
    table = doc.add_table(rows=len(steps) + 1, cols=3)
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    headers = ['Step', 'Action', 'Detail']
    widths  = [Cm(1.5), Cm(4.5), Cm(10)]
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.width = widths[j]
        _bg(c, '008580')
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (num, title, detail) in enumerate(steps):
        bg = 'F0FAFA' if i % 2 == 0 else 'FFFFFF'
        nc = table.cell(i + 1, 0)
        tc = table.cell(i + 1, 1)
        dc = table.cell(i + 1, 2)
        for c, w in zip([nc, tc, dc], widths):
            c.width = w
            _bg(c, bg)
        rn = nc.paragraphs[0].add_run(num)
        rn.bold = True; rn.font.size = Pt(11); rn.font.color.rgb = TEAL; rn.font.name = 'Calibri'
        nc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = tc.paragraphs[0].add_run(title)
        rt.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = NAVY; rt.font.name = 'Calibri'
        rd = dc.paragraphs[0].add_run(detail)
        rd.font.size = Pt(10); rd.font.color.rgb = GREY; rd.font.name = 'Calibri'

    doc.add_paragraph()


def note_box(doc, label, text, colour_hex='FFF8E1', border_hex='C07D00'):
    """Highlighted note / important callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    _bg(cell, colour_hex)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + '  ')
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = AMBER; r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5); r2.font.color.rgb = GREY; r2.font.name = 'Calibri'
    doc.add_paragraph()


# ── Build Document ────────────────────────────────────────────────────────────

doc = Document()

# Page setup — A4
s = doc.sections[0]
s.page_width    = Cm(21.0)
s.page_height   = Cm(29.7)
s.left_margin   = Cm(2.5)
s.right_margin  = Cm(2.5)
s.top_margin    = Cm(2.2)
s.bottom_margin = Cm(2.2)

# Default style
style = doc.styles['Normal']
style.font.name      = 'Calibri'
style.font.size      = Pt(10.5)
style.font.color.rgb = GREY

# ════════════════════════════════════════════════════════════════════
#  COVER
# ════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run('OSCE EXAMINATION MANAGEMENT PLATFORM')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = TEAL; r.font.name = 'Calibri'

doc.add_paragraph()

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run('User Workflow Reference Guide')
r2.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = NAVY; r2.font.name = 'Calibri'

doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = tp.add_run('Examiner Workflow  ·  Coordinator / Creator Workflow')
rt.italic = True; rt.font.size = Pt(11); rt.font.color.rgb = LGREY; rt.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_table(rows=4, cols=2)
meta.style = 'Table Grid'
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_rows = [
    ('Document Type',    'Workflow Reference — Internal Academic Use'),
    ('Intended Readers', 'Examiners, Coordinators, Academic Supervisors'),
    ('Date',             datetime.date.today().strftime('%B %d, %Y')),
    ('Classification',   'CONFIDENTIAL'),
]
for i, (k, v) in enumerate(meta_rows):
    bg = 'F0FAFA' if i % 2 == 0 else 'FFFFFF'
    kc = meta.cell(i, 0); vc = meta.cell(i, 1)
    kc.width = Cm(4.5); vc.width = Cm(11)
    _bg(kc, bg); _bg(vc, bg)
    rk = kc.paragraphs[0].add_run(k)
    rk.bold = True; rk.font.size = Pt(10); rk.font.color.rgb = NAVY; rk.font.name = 'Calibri'
    rv = vc.paragraphs[0].add_run(v)
    rv.font.size = Pt(10); rv.font.color.rgb = GREY; rv.font.name = 'Calibri'

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  PART A — EXAMINER WORKFLOW
# ════════════════════════════════════════════════════════════════════

chapter(doc, 'PART A — EXAMINER WORKFLOW', teal=True)
para(doc,
    'The Examiner is the frontline user of the OSCE platform. Their role is exclusively focused '
    'on scoring students at assigned stations during active examination sessions. The platform '
    'provides examiners with a dedicated, mobile-optimised interface that is entirely separate '
    'from the administrative coordinator panel — ensuring clarity, speed, and security during '
    'live examination conditions.'
)

# ── A1: Login ────────────────────────────────────────────────────────────────
section(doc, 'A1.  Login and First-Time Access')
para(doc,
    'All examiners access the platform through the unified login portal. On first login, the '
    'system enforces a mandatory password change before granting access to any function. '
    'This ensures each examiner holds a personal credential from the moment they enter the system.'
)
bul(doc, [
    ('Login URL:', 'The institution-provided login address — shared by all user roles.'),
    ('Credentials:', 'Username and password assigned by the system administrator. Default password must be changed on first login.'),
    ('Role Detection:', 'After authentication, the system automatically directs the examiner to the examiner interface. No manual navigation is required.'),
    ('Session Timeout:', 'Examiner sessions expire after 30 minutes of inactivity — designed to protect student data on shared tablets during inter-student transitions.'),
])

# ── A2: Home Dashboard ───────────────────────────────────────────────────────
section(doc, 'A2.  Examiner Home Dashboard')
para(doc,
    'The home dashboard is the examiner\'s central navigation point. It is intentionally minimal — '
    'showing only what is relevant to the current examiner at the current moment.'
)
bul(doc, [
    ('Running Stations:', 'Stations the examiner is assigned to within sessions currently in in_progress status. These are displayed as prominent action cards and represent the examiner\'s immediate work.'),
    ('Upcoming Assignments:', 'Stations within scheduled (not yet activated) sessions. These are shown for awareness; scoring is not yet available.'),
    ('Assignment Cards Display:', 'Each card shows the station number, station name, rotation path, number of students to be scored, duration (in minutes), and maximum achievable score.'),
    ('No Cross-Access:', 'Examiners see only their own assignments. No other examiner\'s stations or scores are visible at any point.'),
])
note_box(doc,
    'NOTE:',
    'If a session has not been activated by the coordinator, no action cards appear for that session. '
    'Examiners must wait for coordinator activation before scoring becomes available.'
)

# ── A3: Station Dashboard ────────────────────────────────────────────────────
section(doc, 'A3.  Station Dashboard — Scoring Progress Overview')
para(doc,
    'Tapping an assignment card opens the Station Dashboard — a dedicated view for monitoring '
    'and managing the scoring progress at that station.'
)
bul(doc, [
    ('Student List:', 'All students assigned to the examiner\'s path are listed in sequence order.'),
    ('Scoring Status per Student:', 'Each student row shows whether they have been scored (submitted), are in progress, or are yet to be marked.'),
    ('Dual-Examiner Awareness:', 'If a second examiner is also assigned to the same station, the dashboard reflects whether each student has been scored by the other examiner — without revealing the other examiner\'s individual scores.'),
    ('Progress Indicator:', 'A percentage progress bar shows how many of the station\'s students have been scored, refreshing dynamically as marks are submitted.'),
    ('Remaining Count:', 'A live counter displays the number of students yet to be scored, helping the examiner manage time during the rotation window.'),
])

# ── A4: Marking Interface ────────────────────────────────────────────────────
section(doc, 'A4.  Marking Interface — Scoring a Student')
para(doc,
    'The marking interface is the primary operational screen during a live examination. '
    'It is designed for one-handed tablet operation at an examination station.'
)

subsection(doc, 'A4.1  Opening a Student Record')
bul(doc, [
    'Select a student from the station dashboard or the student selection screen.',
    'The system opens the student\'s marking sheet for that station.',
    'If the examiner has scored this student previously (within this session and station), the previously entered scores are restored automatically — enabling seamless resumption after interruptions.',
])

subsection(doc, 'A4.2  Station Scenario and Instructions Panel')
bul(doc, [
    ('Station Scenario:', 'The clinical scenario for the station is displayed at the top of the marking interface, ensuring the examiner reviews the case context before scoring begins.'),
    ('Examiner Instructions:', 'Any specific scoring instructions authored by the coordinator are surfaced directly in the interface — no paper reference required.'),
    ('Duration Display:', 'The allocated station duration is shown, helping the examiner manage the scoring window.'),
])

subsection(doc, 'A4.3  Checklist-Based Scoring')
para(doc,
    'Each station has a structured checklist of clinical performance items authored during exam '
    'design. The marking interface renders every item with its rubric — interactive scoring '
    'buttons that adapt to the type of assessment required for each item.'
)
bul(doc, [
    ('Binary Items (Pass / Fail):', '"Not Done" and "Done" buttons — tap once to score. Suitable for binary clinical actions (e.g., washes hands, introduces self, obtains consent).'),
    ('Partial Credit Items:', '"Not Done", "Partial", and "Complete" — three-state scoring for tasks where partial performance is clinically meaningful.'),
    ('Scale Items:', 'A numbered progression of score buttons (e.g., 0 through 5) for assessments requiring granular clinical judgement.'),
    ('ILO / Theme Labels:', 'Each checklist item displays its mapped Intended Learning Outcome and clinical theme — contextualising the item within the academic curriculum.'),
    ('Auto-Save:', 'Every tap of a score button immediately persists the score via a background API call. There is no separate save step — the scoring record is continuously updated.'),
    ('Running Total:', 'A live total score counter updates after each item is scored, showing the examiner the student\'s accumulating score against the station maximum.'),
])

subsection(doc, 'A4.4  Global Rating and Comments')
bul(doc, [
    ('Global Rating:', 'An optional overall performance rating field is available at the bottom of the checklist, allowing the examiner to record an holistic clinical impression separate from item-by-item scoring.'),
    ('Examiner Comments:', 'A free-text field accepts examiner remarks that are stored with the score record — accessible to supervisors in the administrative interface.'),
])

subsection(doc, 'A4.5  Submitting the Score')
bul(doc, [
    ('Submit Button:', 'Tapping Submit finalises the score for that student at that station. The score status transitions from in_progress to submitted.'),
    ('Student Status Update:', 'Upon submission, the student\'s session record is automatically updated — moving to in_progress (partial completion) or completed (all stations marked) based on path station count.'),
    ('5-Minute Undo Window:', 'Immediately after submission, a brief window allows the examiner to recall and reopen the score for editing — useful for accidental early submission. After 5 minutes, the score is locked.'),
    ('Post-Submit Review Mode:', 'After the undo window closes, the examiner may re-open the scoring view in read-only review mode — scores are visible but cannot be altered without supervisor-authorised correction.'),
    ('Correction Workflow:', 'If a coordinator unlocks a score for correction, the examiner is able to re-enter the marking interface in edit mode, make amendments, and resubmit. The correction event is written to the audit log with before/after values.'),
])

# ── A5: Dual-Examiner ────────────────────────────────────────────────────────
section(doc, 'A5.  Dual-Examiner Stations')
para(doc,
    'Where two examiners are assigned to the same station — a common clinical validity measure — '
    'the platform manages both scoring streams independently before automatically reconciling them.'
)
bul(doc, [
    'Each examiner scores the same student entirely independently, with no access to the other\'s scores during the session.',
    'Upon both examiners submitting for a student, the station dashboard displays the co-examiner\'s name and their submitted total.',
    'The platform automatically calculates the final score as the arithmetic average of both examiners\' totals.',
    'The averaged final score is used in all coordinator reports and student result calculations.',
])

# ── A6: Offline PWA ──────────────────────────────────────────────────────────
section(doc, 'A6.  Offline Operation and Background Synchronisation')
para(doc,
    'The examiner interface is built as a Progressive Web Application (PWA) — capable of full '
    'operation in environments where network connectivity is absent or unreliable, such as '
    'clinical simulation suites, basement examination venues, or remote facilities.'
)
bul(doc, [
    ('Offline Availability:', 'The examiner interface is cached on the device at first load. If network connectivity is lost, the application continues to function without interruption.'),
    ('Local Score Storage:', 'When offline, all score submissions are queued in the device\'s local storage with a unique identifier (UUID) and local timestamp.'),
    ('Automatic Sync:', 'When network connectivity is restored, the platform\'s Service Worker automatically identifies queued records and synchronises them to the server in the background — without examiner action.'),
    ('Conflict Resolution:', 'The sync engine compares local timestamps with server records. Only records where the local version is newer than the server version are applied. Conflicts are flagged in the system log.'),
    ('Assignment Re-Validation:', 'Every synced record is validated against the examiner\'s assignment data before being persisted — preventing injection of scores outside the examiner\'s authorised stations.'),
    ('Sync Status Indicator:', 'An API endpoint allows the interface to display how many score records are pending synchronisation, providing real-time feedback to the examiner.'),
    ('Offline Fallback Page:', 'A dedicated offline page is served if the examiner navigates to an uncached route while offline, providing a clear degradation experience.'),
])

# ── A7: All Sessions ─────────────────────────────────────────────────────────
section(doc, 'A7.  Session History')
para(doc,
    'The All Sessions view provides a chronological record of all examination sessions to '
    'which the examiner has been assigned — including both active and historical sessions. '
    'This allows examiners to review past assignments and, where permitted by session status, '
    'access previously submitted score records in review mode.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  PART B — CREATOR / COORDINATOR WORKFLOW
# ════════════════════════════════════════════════════════════════════

chapter(doc, 'PART B — COORDINATOR / CREATOR WORKFLOW', teal=True)
para(doc,
    'The Coordinator (also referred to as the Creator) is the primary administrative user '
    'responsible for the full lifecycle of OSCE examination design, session delivery, and '
    'post-session reporting. The coordinator interface provides comprehensive tools for '
    'every stage — from initial course and exam setup through to final mark export.'
)

# ── B1: Login and Dashboard ──────────────────────────────────────────────────
section(doc, 'B1.  Login and Coordinator Dashboard')
bul(doc, [
    ('Login:', 'Coordinators use the same unified login portal as examiners. Role-based redirection delivers them directly to the coordinator dashboard.'),
    ('Dashboard Overview:', 'The coordinator dashboard presents a consolidated summary: total exam count, active sessions currently in progress, draft exams awaiting finalisation, and recently created exams.'),
    ('Quick Navigation:', 'All major functional areas — Courses, Exams, Examiners, Library, Reports — are accessible from the dashboard navigation panel.'),
])

# ── B2: Course Setup ─────────────────────────────────────────────────────────
section(doc, 'B2.  Course and ILO Setup')
para(doc,
    'Before an exam can be created, the course structure must be established. Courses represent '
    'the academic units to which OSCE exams are attached.'
)
bul(doc, [
    ('Create a Course:', 'Define the course code, full name, academic year level (Year-4 Junior or Year-6 Senior), and OSCE mark weighting.'),
    ('Define ILOs:', 'For each course, create the Intended Learning Outcomes that will be mapped to individual checklist items during station design. Each ILO carries a number, description, and maximum OSCE mark allocation.'),
    ('Themes:', 'Group ILOs and stations by clinical theme (e.g., History Taking, Physical Examination, Communication) to support thematic performance analytics.'),
])

# ── B3: Exam Creation ────────────────────────────────────────────────────────
section(doc, 'B3.  Exam Creation')
workflow_table(doc, [
    ('1', 'Navigate to Exams',     'Select "Exams" from the sidebar, then click "Create New Exam".'),
    ('2', 'Link to a Course',      'Select the course this exam belongs to. The exam inherits ILOs and theme structure from the course.'),
    ('3', 'Name the Exam',         'Enter a descriptive exam title (e.g., "Internal Medicine OSCE — Year 6 Semester 2").'),
    ('4', 'Set Exam Weight',       'Enter the OSCE exam weight as a percentage of the total course mark.'),
    ('5', 'Save as Draft',         'The exam is saved in Draft status. Stations and sessions can be added at any time before activation.'),
])

# ── B4: Station Design ───────────────────────────────────────────────────────
section(doc, 'B4.  Station Design')
para(doc,
    'Stations are the individual clinical assessment posts within an exam. Each station has '
    'a clinical scenario, examiner instructions, a duration, and a checklist of performance items.'
)
bul(doc, [
    ('Station Number:', 'A unique number identifying the station within the exam (used for path sequencing).'),
    ('Scenario:', 'The clinical case presented to the student at the station — authored directly in the platform.'),
    ('Instructions:', 'Specific guidance for the examiner on how to administer and observe the station.'),
    ('Duration:', 'The time allocated per student at this station (in minutes).'),
    ('Checklist Items:', 'Individual scored performance items. Each item specifies the description, point value, rubric type (binary / partial / scale), clinical category, and ILO mapping.'),
    ('Critical Item Flag:', 'Individual items may be flagged as critical — enabling fail-override logic if a critical item is missed, regardless of total score.'),
    ('Station Library:', 'Frequently used checklist templates can be saved to the institutional library for reuse across multiple exams, eliminating duplicate authoring effort.'),
])

# ── B5: Session Creation ─────────────────────────────────────────────────────
section(doc, 'B5.  Session Creation and Path Configuration')
para(doc,
    'An exam session represents a scheduled delivery instance of the exam — typically a morning '
    'or afternoon block. Each session operates independently with its own paths, students, '
    'examiner assignments, and scoring data.'
)
workflow_table(doc, [
    ('1', 'Create Session',         'From the exam detail page, click "Add Session". Enter the session name, date, time, and type (Morning / Afternoon).'),
    ('2', 'Configure Paths',        'Specify the number of rotation paths for the session (e.g., 5 paths of 10 students each). The platform creates the path structure automatically.'),
    ('3', 'Assign Stations to Paths','Stations are replicated across all paths. Each path receives an identical station layout, ensuring all students are assessed against the same criteria.'),
    ('4', 'Enrol Students',         'Import a student list (CSV / Excel) or add students manually. Each student is assigned to a path with a sequence number determining their rotation order.'),
    ('5', 'Assign Examiners',       'For each station on each path, assign the attending examiner. One or two examiners may be assigned per station to support dual-examiner protocols.'),
    ('6', 'Review & Activate',      'Review the session configuration. When ready, set the session status to In Progress — this makes it live and accessible to examiners.'),
])
note_box(doc,
    'IMPORTANT:',
    'Examiners cannot see or score any session until the coordinator changes its status to '
    '"In Progress". This ensures no premature scoring occurs before the examination begins.'
)

# ── B6: Live Session Monitoring ───────────────────────────────────────────────
section(doc, 'B6.  Live Session Monitoring')
para(doc,
    'During an active session, the coordinator retains full visibility of scoring progress '
    'across all paths and stations in real time.'
)
bul(doc, [
    ('Session Results Dashboard:', 'Displays each student\'s current score across all stations — updated as examiners submit marks. Sorted by student name with search capability.'),
    ('Per-Student Scoring View:', 'Drill into any student\'s record to see their station-by-station breakdown, total score, percentage, and pass/fail status.'),
    ('Completion Progress:', 'The session summary shows how many students have been fully scored, are in progress, or have not yet been reached.'),
    ('Live Student Search:', 'A real-time search bar on the session view allows immediate student lookup by name or student number — critical during large sessions.'),
    ('Score Unlock (Correction):', 'The coordinator may flag individual station scores for correction. This re-opens the marking interface for the assigned examiner within a controlled, audited correction workflow.'),
])

# ── B7: Examiner Management ───────────────────────────────────────────────────
section(doc, 'B7.  Examiner Management')
bul(doc, [
    ('Create Individual Examiners:', 'Add examiners directly with their name, username, email (optional), and department.'),
    ('Bulk Import:', 'Upload a CSV or Excel file to import multiple examiners at once — suitable for onboarding clinical faculty at the start of an academic semester.'),
    ('Soft Delete:', 'Removed examiners are soft-deleted — their historical scoring records are fully preserved and remain accessible in all prior session reports. They can be restored if required.'),
    ('Password Management:', 'Coordinators can reset examiner passwords from the management interface.'),
])

# ── B8: Reports ───────────────────────────────────────────────────────────────
section(doc, 'B8.  Reporting and Data Export')
para(doc,
    'Once a session is completed, the platform provides a comprehensive suite of reporting '
    'tools — all accessible from the Reports section without requiring manual data assembly.'
)
bul(doc, [
    ('Student Results XLSX:', 'An Excel workbook containing each student\'s station scores, total score, maximum score, percentage, and pass/fail indicator. Formatted for immediate distribution or archival.'),
    ('ILO Score Report XLSX:', 'A per-student, per-ILO performance matrix showing how each student performed against every Intended Learning Outcome across all stations — directly aligned with accreditation evidence requirements.'),
    ('Session Report PDF:', 'A formal A4 PDF document generated via the browser print engine. Contains five sections: Session Overview, Participant Summary, Examiner Assignments, Path & Station Breakdown, and full Student Score Listing. Carries a report code (RPT-XXXX-XXXXXXXX), document date, and a repeating CONFIDENTIAL footer on every page.'),
    ('Raw CSV Export:', 'A raw scoring data export for integration with institutional MIS systems, statistics tools, or custom analysis pipelines.'),
    ('Session Selector:', 'All reports are scoped to a selected session. Coordinators choose the target session from a dropdown before generating any report.'),
])

# ── B9: Session Lifecycle ─────────────────────────────────────────────────────
section(doc, 'B9.  Session Lifecycle — Status Transitions')
workflow_table(doc, [
    ('Draft',       'Exam Not Yet Active',       'Exam exists; stations and sessions are being designed. No examiner access.'),
    ('Scheduled',   'Session Created',           'Session is configured but not yet live. Visible to examiners as "upcoming" but scoring is disabled.'),
    ('In Progress', 'Session Activated',         'Coordinator activates the session. Examiners can now log in and begin marking. All scoring APIs accept submissions.'),
    ('Completed',   'Session Closed',            'Coordinator marks session complete. No further scoring is accepted. Full report suite becomes available.'),
])

# ── B10: Security and Governance ──────────────────────────────────────────────
section(doc, 'B10.  Administrative Controls and Access Governance')
bul(doc, [
    ('Role Restrictions:', 'Coordinators have access to exam management and session operations. Access to system-level administration (user management, audit logs, platform settings) is restricted to administrators and superusers only.'),
    ('Audit Log:', 'Every significant coordinator action — session activation, student enrolment, examiner assignment, score correction — is written to an immutable audit log accessible to administrators.'),
    ('Soft Delete Architecture:', 'Deleted exams, stations, examiners, and sessions are archived rather than destroyed, ensuring historical data integrity and recovery capability.'),
    ('Data Export Control:', 'Report generation and data export functions are limited to authenticated coordinator and administrator roles; examiners have no access to aggregated score data.'),
])

# ── Closing ───────────────────────────────────────────────────────────────────
doc.add_page_break()
_rule(doc, '008580', sz=12)
cp2 = doc.add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = cp2.add_run(
    f'CONFIDENTIAL  ·  OSCE Examination Management Platform  ·  Workflow Reference  ·  '
    f'{datetime.date.today().strftime("%B %d, %Y")}'
)
rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = LGREY; rc.font.name = 'Calibri'

# ── Save ─────────────────────────────────────────────────────────────────────
base     = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
out_path = os.path.join(base, 'OSCE_Workflow_Guide.docx')
doc.save(out_path)
print(f'Saved → {out_path}')
