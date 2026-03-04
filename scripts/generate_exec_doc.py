"""
Generate OSCE Platform – Executive Briefing Document (.docx)
Run: python scripts/generate_exec_doc.py
Output: OSCE_Executive_Briefing.docx (project root)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ──────────────────────────────────────────────────────────
TEAL        = RGBColor(0x00, 0x85, 0x80)   # brand primary
DARK_NAVY   = RGBColor(0x1A, 0x1A, 0x2E)   # headings
MID_GREY    = RGBColor(0x44, 0x44, 0x55)   # body
LIGHT_TEAL  = RGBColor(0xE0, 0xF5, 0xF4)   # table header bg
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '008580')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(0)
    return p


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL
    run.font.name = 'Calibri'
    add_horizontal_rule(doc)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_NAVY
    run.font.name = 'Calibri'
    return p


def body(doc, text, bold=False, italic=False, colour=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    run.font.color.rgb = colour or MID_GREY
    run.font.name = 'Calibri'
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = MID_GREY
    run.font.name = 'Calibri'
    return p


def two_col_table(doc, rows, header=None):
    """Render a clean two-column key-value table."""
    table = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [Cm(5.5), Cm(10.5)]

    if header:
        for j, h in enumerate(header):
            cell = table.cell(0, j)
            cell.width = col_widths[j]
            set_cell_bg(cell, '008580')
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    offset = 1 if header else 0
    for i, (key, val) in enumerate(rows):
        bg = 'F0FAFA' if i % 2 == 0 else 'FFFFFF'
        key_cell = table.cell(i + offset, 0)
        val_cell = table.cell(i + offset, 1)
        key_cell.width = col_widths[0]
        val_cell.width = col_widths[1]
        set_cell_bg(key_cell, bg)
        set_cell_bg(val_cell, bg)
        kr = key_cell.paragraphs[0].add_run(key)
        kr.bold = True; kr.font.size = Pt(10); kr.font.color.rgb = DARK_NAVY
        kr.font.name = 'Calibri'
        vr = val_cell.paragraphs[0].add_run(val)
        vr.font.size = Pt(10); vr.font.color.rgb = MID_GREY
        vr.font.name = 'Calibri'

    doc.add_paragraph()
    return table


# ── Document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── Cover Page ───────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run('OSCE EXAMINATION MANAGEMENT PLATFORM')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = TEAL; r.font.name = 'Calibri'

doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('Executive Technical Briefing')
rs.font.size = Pt(14); rs.italic = True; rs.font.color.rgb = DARK_NAVY; rs.font.name = 'Calibri'

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = tagline.add_run('Prepared for Academic and Institutional Leadership Review')
rt.font.size = Pt(11); rt.font.color.rgb = MID_GREY; rt.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ('Document Classification', 'CONFIDENTIAL – For Internal Academic Use Only'),
    ('Prepared By',             'Platform Development Team'),
    ('Review Audience',         'University Dean, Academic Board, IT Governance'),
    ('Document Date',           datetime.date.today().strftime('%B %d, %Y')),
    ('Platform Version',        'Release 1.0  |  Django 5.2  |  Python 3.13'),
]
two_col_table(doc, meta)

doc.add_page_break()

# ── 1. Executive Summary ────────────────────────────────────────────────────

heading1(doc, '1. Executive Summary')
body(doc,
    'The OSCE Examination Management Platform is an enterprise-grade, purpose-built digital '
    'infrastructure designed to modernise the administration, delivery, and analysis of '
    'Objective Structured Clinical Examinations (OSCE) within academic medical institutions. '
    'Developed on a robust, security-hardened technology foundation, the platform consolidates '
    'the entire OSCE lifecycle — from exam design and session scheduling through real-time '
    'examiner scoring to institutional reporting — into a single, cohesive system.'
)
body(doc,
    'The platform eliminates paper-based workflows, reduces administrative burden, eliminates '
    'data transcription errors, and delivers immediate, auditable scoring outcomes. Its '
    'offline-capable Progressive Web Application (PWA) examiner interface ensures uninterrupted '
    'operation in clinically challenging environments where network connectivity may be limited. '
    'Full Arabic language support ensures inclusive deployment across bilingual academic contexts.'
)
body(doc,
    'This document presents a comprehensive technical and functional review of the platform '
    'for evaluation by academic leadership and institutional governance bodies.'
)

# ── 2. System Overview ───────────────────────────────────────────────────────

heading1(doc, '2. System Overview and Objectives')
body(doc,
    'The platform was architected with four primary institutional objectives:'
)
bullet(doc, 'Standardise OSCE delivery — enforce consistent rubric-based assessment across all examiners and paths simultaneously.')
bullet(doc, 'Eliminate administrative overhead — replace manual paper checklists, score transcription, and spreadsheet aggregation.')
bullet(doc, 'Deliver real-time institutional insight — provide coordinators and academic leadership with live scoring data and immediate post-session analytics.')
bullet(doc, 'Ensure institutional accountability — establish end-to-end audit trails for all scoring events, user actions, and administrative operations.')
doc.add_paragraph()

two_col_table(doc, [
    ('Platform Type',       'Multi-role web application (coordinator + examiner interfaces)'),
    ('Primary Users',       'Exam Coordinators, Examiners, Academic Administrators'),
    ('Languages Supported', 'English and Arabic (full RTL support with HarfBuzz rendering)'),
    ('Deployment Model',    'On-premises or cloud-hosted (Railway / Render / Azure compatible)'),
    ('Architecture',        'Django MVC, REST JSON API, Progressive Web Application (PWA)'),
], header=['Attribute', 'Value'])

# ── 3. Core Features ─────────────────────────────────────────────────────────

heading1(doc, '3. Core Features and Functional Capabilities')

heading2(doc, '3.1  Exam and Course Configuration')
bullet(doc, 'Hierarchical course structure: Courses → Exams → Sessions → Paths → Stations → Checklist Items.')
bullet(doc, 'Intended Learning Outcome (ILO) integration — each checklist item is mapped to a specific ILO for granular academic outcome tracking.')
bullet(doc, 'Support for Year-4 (Junior) and Year-6 (Senior) clinical course tiers.')
bullet(doc, 'Station template library for reusable, validated checklist patterns across multiple exams.')
bullet(doc, 'Rubric configurability per item: binary (pass/fail), partial credit, and multi-level scale types.')
bullet(doc, 'Float-precision scoring supporting clinically precise fractional marks.')
doc.add_paragraph()

heading2(doc, '3.2  Session and Multi-Path Rotation Management')
bullet(doc, 'Session scheduling with configurable start times, session types (morning / afternoon), and status lifecycle (scheduled → in_progress → completed).')
bullet(doc, 'Multi-path circuit management: up to 10 rotation paths per session, with simultaneous parallel execution.')
bullet(doc, 'Automatic student path assignment with redistribution capability via bulk database operations.')
bullet(doc, 'Examiner assignment per station per path per session — enforced at API level.')
bullet(doc, 'Dual-examiner scoring support: when two examiners score the same student at a station, the platform automatically averages both scores to produce a validated final mark.')
doc.add_paragraph()

heading2(doc, '3.3  Examiner Interface')
body(doc,
    'The examiner-facing interface is a mobile-optimised, tablet-ready Progressive Web Application '
    'designed for use in fast-paced clinical examination environments. Its architecture prioritises '
    'speed, reliability, and simplicity of interaction under examination conditions.'
)
bullet(doc, 'Personalised dashboard — each examiner sees only their assigned stations and active sessions; no access to other examiners\' data.')
bullet(doc, 'Session status enforcement — the station dashboard is accessible only when the session status is in_progress; coordinators control activation.')
bullet(doc, 'Station assignment cards — display station number, name, path, student count, duration, and maximum score at a glance.')
bullet(doc, 'Checklist-driven scoring interface — each item is presented with contextual rubric options rendered as tap-friendly score buttons.')
bullet(doc, 'Rubric type rendering — binary items show "Done / Not Done" buttons; partial items show "Not Done / Partial / Complete"; scale items show a numbered progression.')
bullet(doc, 'Scenario and instructions panel — station scenario and examiner instructions are surfaced directly within the scoring view.')
bullet(doc, 'ILO tagging — each checklist item displays its mapped ILO and theme to inform examiner clinical judgement.')
bullet(doc, 'Resume capability — if an examiner exits mid-scoring, the system restores their previous in-progress entries on re-entry.')
bullet(doc, 'Dual-examiner awareness — the dashboard shows whether a co-examiner has already submitted a score for a student.')
bullet(doc, 'Offline-first PWA with background synchronisation — scoring continues without network connectivity; all data is queued locally and synced automatically when connectivity is restored.')
bullet(doc, 'Activity-based session timeout (30 minutes inactivity) protecting patient and student data on shared tablets.')
bullet(doc, 'Dedicated offline fallback page ensuring graceful degradation in zero-connectivity scenarios.')
bullet(doc, 'Unified login with role-based redirection — examiners are directed exclusively to their interface upon authentication.')
bullet(doc, 'All-sessions history view — examiners can review their full assignment history across past sessions.')
doc.add_paragraph()

heading2(doc, '3.4  Coordinator and Administrative Interface')
bullet(doc, 'Full exam lifecycle management: create, edit, activate, complete, and archive exams and sessions.')
bullet(doc, 'Real-time session monitoring: live scoring progress dashboard per student per station.')
bullet(doc, 'Examiner management: create accounts, assign to stations, bulk import via CSV/Excel.')
bullet(doc, 'Student enrolment and path assignment with bulk redistribution.')
bullet(doc, 'Soft-delete with restore capability for all critical entities.')
doc.add_paragraph()

heading2(doc, '3.5  Reporting and Analytics')
bullet(doc, 'Session Results Report: per-student total score, max score, percentage, pass/fail indicator with station breakdown.')
bullet(doc, 'ILO Score Export: per-student, per-ILO score matrix exported to formatted XLSX — directly usable for accreditation reporting.')
bullet(doc, 'Student Results XLSX: formatted Excel workbook with per-station scores and summary statistics.')
bullet(doc, 'Session Report PDF: A4 print-ready document with five structured sections, report code, page numbers, and CONFIDENTIAL footer — suitable for official institutional records.')
bullet(doc, 'Raw CSV export with optimised database queries for integration with institutional data warehouses.')
doc.add_paragraph()

# ── 4. User Experience ───────────────────────────────────────────────────────

heading1(doc, '4. User Experience and Interface Design')
body(doc,
    'The platform employs a deliberate split-interface architecture: a full-featured coordinator '
    'dashboard built on Bootstrap 5 with a professional dark-navy / teal design language, and a '
    'lightweight, high-contrast examiner interface optimised for tablet and mobile use under '
    'time-pressured examination conditions.'
)
bullet(doc, 'Responsive layouts — fully adaptive across desktop, tablet, and mobile viewport sizes.')
bullet(doc, 'Role-isolated navigation — each user role sees only the interface and data relevant to their function.')
bullet(doc, 'Contextual in-page feedback — success, warning, and error messages delivered without full page reloads.')
bullet(doc, 'Force-change-password workflow — first-login password update enforced before system access is granted.')
bullet(doc, 'Arabic language support — full RTL text rendering using browser-native HarfBuzz / DirectWrite engines; Arabic student names preserved correctly in all reports.')
bullet(doc, 'Progressive Web Application — the examiner interface is installable on any device as a native-like app, supporting home-screen launch and offline access.')
doc.add_paragraph()

# ── 5. Security Architecture ─────────────────────────────────────────────────

heading1(doc, '5. Security Architecture')
body(doc,
    'Security is treated as a foundational engineering requirement, not an afterthought. The '
    'platform has undergone a comprehensive security audit resulting in nine discrete hardening '
    'measures. The following controls are active in the production-configured system:'
)

heading2(doc, '5.1  Authentication and Access Control')
bullet(doc, 'Four-tier role hierarchy: Superuser → Admin → Coordinator → Examiner — enforced at middleware, view, and API levels.')
bullet(doc, 'django-axes integration — automated brute-force protection with configurable lockout thresholds and IP-based rate limiting on the login endpoint.')
bullet(doc, 'Activity-based session timeouts differentiated by role (5 minutes for administrative roles; 30 minutes for examiners).')
bullet(doc, 'Django admin portal restricted exclusively to Superuser accounts; hidden from all other role levels.')
bullet(doc, 'Secret admin URL token scoped only to authenticated staff users — never exposed in examiner-facing templates.')
doc.add_paragraph()

heading2(doc, '5.2  API and Endpoint Protection')
bullet(doc, 'Middleware-enforced API path guard — all /api/creator/ endpoints are blocked for non-coordinator roles at the infrastructure layer.')
bullet(doc, 'Examiner assignment verification — all scoring endpoints (get_session_students, start_marking, sync_offline_data) verify the requesting examiner holds a valid assignment before processing.')
bullet(doc, 'Session status validation — start_marking rejects requests against completed or cancelled sessions, preventing unauthorised score backdating.')
bullet(doc, 'Ownership checks on sync — sync_offline_data verifies the examiner owns the score record before permitting updates.')
bullet(doc, 'JSON body validation — all 13+ JSON parsing operations are wrapped in structured error handlers returning HTTP 400 on malformed input.')
doc.add_paragraph()

heading2(doc, '5.3  Data Integrity and CSRF Protection')
bullet(doc, 'Django CSRF middleware active on all form submissions and AJAX scoring calls; unnecessary exemptions removed during audit.')
bullet(doc, 'Soft-delete architecture — critical academic records (stations, scores, examiners) are never physically destroyed; audit trail is preserved.')
bullet(doc, 'Comprehensive audit logging — all critical actions (login, logout, score creation, admin operations) are written to an immutable audit log with actor, timestamp, and context.')
doc.add_paragraph()

heading2(doc, '5.4  Compliance Considerations')
bullet(doc, 'All credentials and secret keys managed via environment variables (.env); no hardcoded secrets in source code.')
bullet(doc, 'Production deployment enforces DEBUG=False, strict ALLOWED_HOSTS, and HTTPS-ready configuration.')
bullet(doc, 'Data handling aligned with institutional data governance requirements; student records are access-controlled by role and session assignment.')
doc.add_paragraph()

# ── 6. Performance and Scalability ──────────────────────────────────────────

heading1(doc, '6. System Performance and Scalability')
body(doc,
    'The platform has undergone a dedicated performance optimisation review addressing seven '
    'identified bottlenecks. The system is designed to sustain concurrent multi-session '
    'examination delivery without degradation in response time or data consistency.'
)

two_col_table(doc, [
    ('N+1 Query Elimination',   'select_related() and prefetch_related() applied across all major API endpoints; annotate() used for aggregate counts, eliminating per-row database calls.'),
    ('Bulk Operations',         'Student redistribution uses Django bulk_update() — single database round-trip regardless of student count.'),
    ('Database Indexing',       'db_index=True applied to ExamSession.status and SessionStudent.status — critical fields used in high-frequency filtering queries.'),
    ('Paginated API Responses', 'All large result sets (student lists, score summaries) are paginated with configurable page size to bound response payload.'),
    ('Offline Sync Architecture', 'PWA queues scores locally; background sync batches updates — server load is smoothed across the session timeline rather than spiking at completion.'),
    ('Static File Serving',     'WhiteNoise middleware serves compressed, cache-headered static assets with zero additional infrastructure.'),
], header=['Optimisation', 'Implementation Detail'])

body(doc,
    'For institutional-scale deployments (500+ concurrent students across multiple sessions), '
    'the architecture supports a straightforward migration from SQLite to PostgreSQL via a '
    'single environment variable change, and horizontal scaling via Gunicorn multi-worker '
    'configuration with a reverse proxy (nginx / Azure Application Gateway).'
)

# ── 7. Data Management ───────────────────────────────────────────────────────

heading1(doc, '7. Data Management and Backup Strategy')
bullet(doc, 'Production database: PostgreSQL (via psycopg2) — ACID-compliant, supporting full transactional integrity across concurrent scoring operations.')
bullet(doc, 'Django migration framework ensures schema evolution is version-controlled, reversible, and auditable.')
bullet(doc, 'Soft-delete across critical models — exam scores, student records, and station configurations are logically archived rather than permanently deleted, enabling data recovery and forensic audit.')
bullet(doc, 'Export pipeline — all institutional data is exportable on demand to XLSX and CSV formats, providing a complete institutional backup independent of the database layer.')
bullet(doc, 'Cloud hosting providers (Railway, Render, Azure) provide automated daily database backups with point-in-time recovery.')
bullet(doc, 'Environment-variable-driven configuration ensures complete separation between development, staging, and production data environments.')

# ── 8. Technology Stack ──────────────────────────────────────────────────────

heading1(doc, '8. Technology Stack')
two_col_table(doc, [
    ('Backend Framework',     'Django 5.2.11 (Python 3.13) — enterprise-proven, LTS-aligned'),
    ('Primary Database',      'PostgreSQL (production) / SQLite (development)'),
    ('Authentication',        'Custom Examiner user model + django-axes brute-force protection'),
    ('Static File Serving',   'WhiteNoise — compressed, cache-optimised, no CDN dependency'),
    ('Excel Generation',      'openpyxl 3.1.5 — server-side XLSX with full formatting'),
    ('PDF / Print',           'Browser-native HTML print (HarfBuzz / DirectWrite — full Arabic)'),
    ('Arabic Text Processing','arabic-reshaper + python-bidi (server-side pre-processing)'),
    ('PWA / Offline Sync',    'Service Worker + IndexedDB background sync'),
    ('Application Server',    'Gunicorn (WSGI) — multi-worker production configuration'),
    ('Configuration',         'django-environ — .env-driven, 12-factor compliant'),
    ('Deployment Targets',    'Railway, Render, Azure App Service, on-premises Linux'),
], header=['Component', 'Technology'])

# ── 9. Future Expansion ──────────────────────────────────────────────────────

heading1(doc, '9. Future Expansion and Sustainability')
body(doc,
    'The platform is engineered with a modular, extensible architecture that accommodates '
    'institutional growth without requiring foundational redesign. The following expansion '
    'pathways have been identified and are architecturally supported:'
)
bullet(doc, 'LMS Integration — REST API endpoints can be extended to support bidirectional data exchange with Moodle, Blackboard, or institutional SIS platforms via OAuth 2.0.')
bullet(doc, 'Multi-Faculty Deployment — the course and exam model supports independent faculty namespacing, enabling a single platform instance to serve multiple medical disciplines.')
bullet(doc, 'Video-Assisted OSCE — the station model is extensible to support multimedia scenario delivery (video briefings, imaging attachments) without schema migration.')
bullet(doc, 'Advanced Analytics Dashboard — the scoring data model supports longitudinal student performance tracking, cohort benchmarking, and ILO attainment visualisation as a next-phase addition.')
bullet(doc, 'Mobile Examiner App — the existing PWA architecture provides a direct upgrade path to a native mobile application (React Native / Flutter) consuming the same REST API.')
bullet(doc, 'Automated Reporting — scheduled report generation (post-session XLSX/PDF auto-delivery to coordinators) is achievable via Django-Q or Celery task queues.')
doc.add_paragraph()

# ── 10. Strategic Value ──────────────────────────────────────────────────────

heading1(doc, '10. Strategic Value to the Institution')
body(doc,
    'The OSCE Examination Management Platform delivers measurable institutional value across '
    'four strategic dimensions:'
)

heading2(doc, 'Academic Quality Assurance')
body(doc,
    'By enforcing standardised, rubric-driven assessment across all examiners and paths '
    'simultaneously, the platform directly addresses inter-examiner variability — a '
    'well-documented threat to OSCE validity. ILO-linked scoring provides granular academic '
    'outcome data directly usable for accreditation evidence packages (e.g., NCAAA, WFME, LCME).'
)

heading2(doc, 'Operational Efficiency')
body(doc,
    'Elimination of paper checklists, manual score transcription, and spreadsheet aggregation '
    'reduces administrative processing time per session from hours to minutes. '
    'Post-session reports are generated on demand, with no manual data entry required.'
)

heading2(doc, 'Institutional Risk Management')
body(doc,
    'End-to-end audit logging, role-based access control, and immutable soft-delete records '
    'provide a defensible, legally credible record of all examination events. The platform '
    'eliminates the documentation vulnerabilities inherent in paper-based systems.'
)

heading2(doc, 'Technology Infrastructure Investment')
body(doc,
    'Built on the internationally adopted Django framework with a PostgreSQL backend, the '
    'platform leverages open-source, vendor-independent technology. There are no proprietary '
    'licensing costs, no single-vendor lock-in, and the codebase is fully owned and '
    'maintainable by the institution\'s IT department. The modular architecture ensures '
    'long-term sustainability and clear upgrade pathways aligned with Django\'s LTS roadmap.'
)

# ── Footer note ──────────────────────────────────────────────────────────────

doc.add_page_break()
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = closing.add_run(
    'This document is prepared exclusively for internal academic and institutional review. '
    'Distribution outside the intended audience requires prior written authorisation.'
)
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = MID_GREY
r.font.name = 'Calibri'

# ── Save ─────────────────────────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(__file__), '..', 'OSCE_Executive_Briefing.docx')
doc.save(out_path)
print(f'Document saved → {os.path.abspath(out_path)}')
