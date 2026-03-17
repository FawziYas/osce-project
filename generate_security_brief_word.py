"""
Generates a Word (.docx) document covering all security details built into the OSCE app.
Run: python generate_security_brief_word.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ──────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE
# ──────────────────────────────────────────────────────────────────────────────
DEEP_BLUE   = RGBColor(0x0A, 0x25, 0x40)
ACCENT_BLUE = RGBColor(0x1A, 0x73, 0xE8)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT   = RGBColor(0x20, 0x21, 0x24)
MID_GREY    = RGBColor(0x5F, 0x63, 0x68)
LIGHT_BLUE_BG  = RGBColor(0xE8, 0xF0, 0xFE)
LIGHT_GREEN_BG = RGBColor(0xE6, 0xF4, 0xEA)
GREEN       = RGBColor(0x18, 0x80, 0x38)
RED_BG      = RGBColor(0xFC, 0xE8, 0xE6)
RED_TEXT    = RGBColor(0xB3, 0x14, 0x12)
AMBER_BG    = RGBColor(0xFE, 0xF7, 0xE0)
AMBER_TEXT  = RGBColor(0xF2, 0x99, 0x00)
TABLE_ALT   = RGBColor(0xF8, 0xF9, 0xFA)

# ──────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ──────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (hex string like '0A2540')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_vertical_align(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def add_run_colored(para, text, bold=False, color=None, size=10):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def section_heading(doc, text, level=1):
    """Add a styled section heading."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = DEEP_BLUE if level == 1 else ACCENT_BLUE
    # Bottom border on level-1 headings
    if level == 1:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '0A2540')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return para


def bullet_item(doc, label, rest='', indent=0.4):
    """Add a bullet point. label is bold, rest is normal."""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(indent)
    para.paragraph_format.space_after = Pt(2)
    if label:
        r = para.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = DARK_TEXT
    if rest:
        r2 = para.add_run('  ' + rest)
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK_TEXT
    return para


def sub_bullet(doc, text, indent=1.0):
    para = doc.add_paragraph(style='List Bullet 2')
    para.paragraph_format.left_indent = Cm(indent)
    para.paragraph_format.space_after = Pt(1)
    r = para.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = MID_GREY
    return para


def styled_table(doc, headers, rows, col_widths,
                 header_bg='0A2540', alt_bg='F8F9FA'):
    """Build a table with styled headers and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = Cm(col_widths[i])
        set_cell_bg(cell, header_bg)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else alt_bg
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Cm(col_widths[ci])
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Allow bold label: if tuple passed, (bold_part, normal_part)
            if isinstance(cell_text, tuple):
                r1 = para.add_run(cell_text[0])
                r1.bold = True
                r1.font.size = Pt(9.5)
                r1.font.color.rgb = DARK_TEXT
                if len(cell_text) > 1:
                    r2 = para.add_run('  ' + cell_text[1])
                    r2.font.size = Pt(9.5)
                    r2.font.color.rgb = DARK_TEXT
            else:
                run = para.add_run(cell_text)
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK_TEXT
    return table


def add_cover(doc):
    """Title block."""
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_para.add_run('OSCE Examination Platform')
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = DEEP_BLUE

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(2)
    sr = sub_para.add_run('Security Architecture — Technical Reference Document')
    sr.font.size = Pt(13)
    sr.font.color.rgb = ACCENT_BLUE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(12)
    mr = meta.add_run('Prepared for IT Department Review  |  March 2026  |  Confidential')
    mr.font.size = Pt(9)
    mr.font.color.rgb = MID_GREY
    mr.italic = True

    # Horizontal rule (paragraph border)
    hr = doc.add_paragraph()
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0A2540')
    pBdr.append(bottom)
    pPr.append(pBdr)
    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENT BUILD
# ──────────────────────────────────────────────────────────────────────────────
def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    add_cover(doc)

    # ── 1. OVERVIEW ───────────────────────────────────────────────────────────
    section_heading(doc, '1. Security Overview')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        'The OSCE platform implements security in multiple independent layers — '
        'if one layer is bypassed, the next still protects data. '
        'The design follows the OWASP Top 10 framework and was independently audited in March 2026. '
        'All critical and high-severity findings were resolved before production deployment.'
    )
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_TEXT

    styled_table(doc,
        ['Security Layer', 'Mechanism', 'Protects Against'],
        [
            ('Authentication',        'Session cookies, CSRF, brute-force lockout',    'Unauthorized access, credential stuffing'),
            ('Authorization',         'Role-based middleware + PostgreSQL RLS',         'Privilege escalation, cross-department data access'),
            ('Input Handling',        'Django ORM — 100% parameterized queries',        'SQL Injection (OWASP A03)'),
            ('Transport Security',    'HTTPS/TLS enforced, HSTS 1 year',               'Man-in-the-middle attacks, data interception'),
            ('Browser Security',      'CSP, X-Frame-Options, nosniff, Referrer-Policy','XSS, clickjacking, MIME sniffing'),
            ('Admin Protection',      'Obscured URL + session gate (double-lock)',       'Unauthorized admin access, scanner discovery'),
            ('Audit Trail',           'Full login, access, and data-change logging',    'Non-repudiation, insider threats, incident response'),
            ('Rate Limiting',         'django-axes (IP + username lockout)',            'Brute-force, automated attacks'),
            ('Password Security',     'Policy enforcement + force-change on first login','Weak credentials, default password abuse'),
            ('Session Management',    'Auto-timeout (10/20 min), HttpOnly, SameSite',  'Session hijacking, CSRF, abandoned sessions'),
        ],
        col_widths=[4.5, 6.0, 6.0]
    )
    doc.add_paragraph()

    # ── 2. AUTHENTICATION ─────────────────────────────────────────────────────
    section_heading(doc, '2. Authentication')
    section_heading(doc, '2.1  Session-Based Login', level=2)
    bullet_item(doc, 'Session cookies only —', 'no JWT tokens in URLs or localStorage; tokens cannot be stolen via URL history or referrer leaks')
    bullet_item(doc, 'HttpOnly flag:', 'session cookie is inaccessible to JavaScript — XSS cannot steal it')
    bullet_item(doc, 'SameSite=Lax:', 'cookie is not sent on cross-site requests — prevents CSRF token theft')
    bullet_item(doc, 'Secure flag (production):', 'cookie is only transmitted over HTTPS — never in plaintext')
    bullet_item(doc, 'CSRF token:', 'every POST/PUT/DELETE requires a server-issued CSRF token — blocks cross-site request forgery')

    section_heading(doc, '2.2  Password Policy', level=2)
    bullet_item(doc, 'Minimum 8 characters')
    bullet_item(doc, 'Rejects common passwords', '(Django CommonPasswordValidator against 20,000+ known weak passwords)')
    bullet_item(doc, 'Rejects numeric-only passwords', '(e.g., "12345678" is blocked)')
    bullet_item(doc, 'Rejects username-similar passwords')
    bullet_item(doc, 'Force-change on first login:', 'every new account is issued a temporary password; the user is redirected to a change-password page and CANNOT use any other feature until they change it — enforced at the middleware level, not optional')
    bullet_item(doc, 'Force-change is blocked for API endpoints:', 'API calls return HTTP 403 JSON while password is still temporary — the system cannot be bypassed by calling APIs directly')

    section_heading(doc, '2.3  Brute-Force / Rate Limiting (django-axes 8.3)', level=2)
    bullet_item(doc, 'Account lockout', 'after configurable number of failed attempts — tracks by both username AND IP address')
    bullet_item(doc, 'Cooldown period:', 'locked account cannot attempt login for a defined time window — prevents automated credential stuffing')
    bullet_item(doc, 'Axes middleware position:', 'sits after AuthenticationMiddleware — every request is inspected before reaching any view')
    bullet_item(doc, 'Admin notification:', 'lockout events are logged and visible to administrators')

    section_heading(doc, '2.4  Soft-Delete Protection', level=2)
    bullet_item(doc, 'Deactivated accounts (is_deleted=True) CANNOT authenticate —', 'an explicit check runs after Django\'s own authenticate() call and rejects the login before a session is created')
    bullet_item(doc, 'Vulnerability patched (CRITICAL):', 'the original code relied solely on Django\'s authenticate() which does not check custom soft-delete flags — this was audited and fixed')
    doc.add_paragraph()

    # ── 3. AUTHORISATION / ACCESS CONTROL ────────────────────────────────────
    section_heading(doc, '3. Authorisation & Access Control')
    section_heading(doc, '3.1  Role-Based Access Control (RBAC)', level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Six distinct roles with strictly scoped permissions:')
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_TEXT

    styled_table(doc,
        ['Role', 'Scope', 'Key Restrictions'],
        [
            ('Superuser',           'Global',       'Full access; session delete/revert only at this level'),
            ('Admin',               'Global',       'Manage all departments, exams, sessions, users'),
            ('Coordinator — Head',  'Department',   'Own department only; can view student list; open dry grading'),
            ('Coordinator — Organizer', 'Department','Own department only; manage sessions/exams; open dry grading'),
            ('Coordinator — RTA',   'Department',   'Own department only; real-time access coordination'),
            ('Examiner',            'Station only', 'ONLY their assigned station — cannot see any other station\'s data'),
        ],
        col_widths=[4.5, 3.5, 8.5]
    )
    doc.add_paragraph()

    section_heading(doc, '3.2  Application-Level Enforcement', level=2)
    bullet_item(doc, 'RoleBasedAccessMiddleware:', 'every incoming request is inspected — the user\'s role is verified before the view code even runs')
    bullet_item(doc, 'Custom DRF permission classes:', 'every REST API endpoint declares explicit role requirements (IsAdmin, IsCoordinator, IsExaminer) — no endpoint is accidentally left open')
    bullet_item(doc, 'scope_queryset() helper:', 'all database queries are automatically filtered to the user\'s permitted scope — a coordinator\'s query physically cannot return rows from a different department')
    bullet_item(doc, 'Object-level permission checks:', 'every object lookup (get_object_or_404) includes scope filtering — an IDOR attack returns 404, not data from another department')
    bullet_item(doc, 'Patched CRITICAL vulnerability:', '45+ API endpoints previously had zero department scoping — all fixed (VULN-002, March 2026 audit)')

    section_heading(doc, '3.3  PostgreSQL Row-Level Security (RLS)', level=2)
    bullet_item(doc, 'Database-level guardrail:', 'even if application code had a bug, PostgreSQL itself blocks rows outside the user\'s permitted scope from being returned')
    bullet_item(doc, 'Session variable injection:', 'Django middleware injects 4 variables per request into the PostgreSQL session: app.current_role, app.current_user_id, app.department_id, app.station_ids')
    bullet_item(doc, 'Transaction-scoped:', 'variables are set with set_config(..., true) — they reset automatically after each request; no cross-request data leakage')
    bullet_item(doc, 'RLS policies cover:', 'all core tables (exam sessions, stations, scores, student data, paths, checklist items, item scores)')
    bullet_item(doc, 'PostgreSQL helper functions:', 'app_role(), app_dept_id(), app_station_ids() — RLS policies use only these helpers, never inline SQL expressions')
    bullet_item(doc, 'Examiner granularity:', 'RLS checks the exact station UUID — an examiner\'s DB connection cannot read scores from any other station, even within the same exam')
    doc.add_paragraph()

    # ── 4. INJECTION PROTECTION ───────────────────────────────────────────────
    section_heading(doc, '4. SQL Injection Protection')
    bullet_item(doc, 'Zero raw SQL in the entire codebase —', 'confirmed across all source files in the March 2026 audit')
    bullet_item(doc, 'Django ORM exclusively:', 'all database access goes through the Django ORM which generates parameterized queries — data values are never interpolated into SQL strings')
    bullet_item(doc, 'No .raw() calls:', 'the codebase contains zero instances of QuerySet.raw(), .extra(), or cursor.execute() with string formatting')
    bullet_item(doc, 'ORM parameter binding:', 'even complex filters (date ranges, multi-condition queries, JSON fields) use ORM Q objects and keyword arguments — not string-built SQL')
    bullet_item(doc, 'SQL injection is structurally impossible', 'by the chosen architecture — it does not rely on developer discipline alone')
    doc.add_paragraph()

    # ── 5. TRANSPORT SECURITY ─────────────────────────────────────────────────
    section_heading(doc, '5. Transport Security (HTTPS / TLS)')
    bullet_item(doc, 'HTTPS enforced in production:', 'SECURE_SSL_REDIRECT=True — all HTTP traffic is permanently redirected to HTTPS (HTTP 301)')
    bullet_item(doc, 'HSTS (HTTP Strict Transport Security):', '1-year duration with includeSubDomains — browsers remember to always use HTTPS; cannot be downgraded by attackers')
    bullet_item(doc, 'Database connections encrypted:', 'PostgreSQL sslmode=require — all data between app server and database is encrypted in transit')
    bullet_item(doc, 'Secure cookie flag:', 'SESSION_COOKIE_SECURE=True and CSRF_COOKIE_SECURE=True in production — cookies never sent over plaintext HTTP')
    doc.add_paragraph()

    # ── 6. HTTP SECURITY HEADERS ──────────────────────────────────────────────
    section_heading(doc, '6. HTTP Security Headers')
    p = doc.add_paragraph()
    r = p.add_run('All headers are applied by custom middleware (ContentSecurityPolicyMiddleware, ReferrerPolicyMiddleware, PermissionsPolicyMiddleware) on EVERY response — not just selected pages.')
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after = Pt(6)

    styled_table(doc,
        ['HTTP Header', 'Value', 'Threat Mitigated'],
        [
            ('Content-Security-Policy',     'Restricts script/style/image sources to trusted origins',      'Cross-Site Scripting (XSS)'),
            ('X-Frame-Options',             'DENY',                                                          'Clickjacking (embedding in iframes)'),
            ('X-Content-Type-Options',      'nosniff',                                                       'MIME-type sniffing attacks'),
            ('Referrer-Policy',             'strict-origin-when-cross-origin',                               'URL leakage to third parties'),
            ('Permissions-Policy',          'camera=(), microphone=(), geolocation=()',                      'Unauthorized hardware access'),
            ('X-Robots-Tag',                'noindex, nofollow',                                             'Search engine indexing of sensitive URLs'),
            ('Strict-Transport-Security',   'max-age=31536000; includeSubDomains',                          'Protocol downgrade / MITM'),
            ('CSRF Token (header)',         'X-CSRFToken required on all state-changing requests',           'Cross-Site Request Forgery'),
        ],
        col_widths=[4.5, 5.5, 6.5]
    )
    doc.add_paragraph()

    # ── 7. ADMIN PANEL DOUBLE-LOCK ────────────────────────────────────────────
    section_heading(doc, '7. Admin Panel — Double-Lock Protection')
    bullet_item(doc, 'Standard /admin/ URL returns HTTP 404 —', 'scanner tools and attackers probing the default Django admin path find nothing; it does not exist')
    bullet_item(doc, 'Secret admin URL:', 'the real admin panel is accessed via a configurable secret path (set as SECRET_ADMIN_URL environment variable); changed with zero code modifications')
    bullet_item(doc, 'Session gate (Lock 2):', 'even knowing the secret URL is not enough — a session token (admin_unlocked=True) must be set by successfully POSTing to the admin gateway; direct URL access without the token returns 404')
    bullet_item(doc, 'Gateway requirements:', 'the gateway requires a valid staff login AND the correct secret token — two independent checks')
    bullet_item(doc, 'Staff flag required:', 'user must have is_staff=True; examiners and coordinators cannot access the admin panel regardless of URL knowledge')
    bullet_item(doc, 'No hardcoded URLs:', 'admin URL is never written into templates or JavaScript — it is constructed server-side only')
    doc.add_paragraph()

    # ── 8. SESSION MANAGEMENT ─────────────────────────────────────────────────
    section_heading(doc, '8. Session Management & Timeout')
    bullet_item(doc, 'Activity-based sliding timeout:', 'session expiry is reset on every request — only genuine inactivity triggers logout')
    bullet_item(doc, 'Admin / Coordinator timeout:', '10 minutes of inactivity → automatic logout')
    bullet_item(doc, 'Examiner timeout:', '20 minutes of inactivity → automatic logout (longer to accommodate natural exam pacing)')
    bullet_item(doc, 'SessionTimeoutMiddleware:', 'custom middleware enforces per-role timeouts; Django\'s built-in SESSION_COOKIE_AGE is only a fallback')
    bullet_item(doc, 'Session isolation:', 'Django session backend stores session data server-side; clients hold only an opaque session ID')
    bullet_item(doc, 'Session regeneration:', 'a new session ID is issued on every login — prevents session fixation attacks')
    doc.add_paragraph()

    # ── 9. AUDIT LOGGING ─────────────────────────────────────────────────────
    section_heading(doc, '9. Comprehensive Audit Logging')
    bullet_item(doc, 'Every login attempt logged:', 'timestamp, username, IP address, user-agent, success/failure')
    bullet_item(doc, 'Every HTTP 401 / 403 / 404 response logged:', 'UnauthorizedAccessMiddleware records unauthorized access attempts with full context')
    bullet_item(doc, 'Score submissions:', 'each score write is attributed to a specific examiner, station, and timestamp — full traceability')
    bullet_item(doc, 'Exam session lifecycle:', 'creation, activation, deactivation, and finalization all logged with acting user')
    bullet_item(doc, 'Dry grading modifications:', 'any score adjustment during dry grading review is logged separately')
    bullet_item(doc, 'Admin actions:', 'Django admin panel records all model-level changes automatically')
    bullet_item(doc, 'Asynchronous logging via Celery:', 'audit log writes are queued as background tasks — exam-room response times are not affected by logging overhead')
    bullet_item(doc, 'Non-repudiation guarantee:', 'every data change is attributed to a named, authenticated user — no anonymous modifications are possible anywhere in the system')
    doc.add_paragraph()

    # ── 10. SECURITY AUDIT RESULTS ───────────────────────────────────────────
    section_heading(doc, '10. Security Audit Results — March 2026')
    p = doc.add_paragraph()
    r = p.add_run('A full OWASP Top 10 audit was conducted covering SQL injection, broken access control, IDOR, authentication bypass, and data leakage. Scope: all views, all API endpoints, all middleware.')
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_TEXT
    p.paragraph_format.space_after = Pt(6)

    styled_table(doc,
        ['Severity', 'Found', 'Fixed', 'Remaining'],
        [
            ('CRITICAL', '2', '2', '0 — All resolved'),
            ('HIGH',     '5', '5', '0 — All resolved'),
            ('MEDIUM',   '6', '6', '0 — All resolved'),
            ('LOW',      '3', '0', '3 — Accepted risk (documented, non-critical)'),
            ('TOTAL',    '16','13', '3'),
        ],
        col_widths=[4.0, 2.5, 2.5, 7.5]
    )
    doc.add_paragraph()

    section_heading(doc, '10.1  Critical Findings (Fixed)', level=2)
    bullet_item(doc, 'VULN-001 — Authentication Bypass (CRITICAL):',
                'Soft-deleted accounts could still log in. Fix: explicit is_deleted check added after authenticate(). Users deactivated by admin are permanently blocked.')
    bullet_item(doc, 'VULN-002 — Cross-Department Data Exposure (CRITICAL):',
                '45+ API endpoints had zero department scoping — any coordinator could read/modify data from all departments. Fix: scope_queryset() applied to every endpoint across 10 API files.')

    section_heading(doc, '10.2  High Findings (Fixed)', level=2)
    bullet_item(doc, 'VULN-003:', 'Examiner list not scoped to department — any coordinator could enumerate examiners from all departments')
    bullet_item(doc, 'VULN-004:', 'Student detail endpoint (creator views) returned data without department check')
    bullet_item(doc, 'VULN-005:', 'Report export endpoints allowed cross-department data extraction')
    bullet_item(doc, 'VULN-006:', 'Path management endpoints missing scope filtering — paths from other departments were accessible')
    bullet_item(doc, 'VULN-007:', 'Checklist library endpoints not department-scoped')

    section_heading(doc, '10.3  OWASP Top 10 Coverage', level=2)
    styled_table(doc,
        ['OWASP Category', 'Status'],
        [
            ('A01 — Broken Access Control',                 '✓ Fixed — RLS + scope_queryset + middleware'),
            ('A02 — Cryptographic Failures',                '✓ Addressed — HTTPS, HSTS, SSL DB connections, HttpOnly cookies'),
            ('A03 — Injection (SQL, XSS, Command)',         '✓ SQLi impossible (ORM); XSS mitigated (CSP + Django auto-escaping)'),
            ('A04 — Insecure Design',                       '✓ Addressed — threat modelling, multi-layer defence'),
            ('A05 — Security Misconfiguration',             '✓ Addressed — DEBUG=False, no default admin URL, secret keys in env vars'),
            ('A06 — Vulnerable & Outdated Components',      '✓ All dependencies current as of March 2026'),
            ('A07 — Identification & Auth Failures',        '✓ Fixed — brute-force lockout, force-change, session management'),
            ('A08 — Software & Data Integrity Failures',    '✓ Addressed — migration audit trail, no untrusted deserialisation'),
            ('A09 — Security Logging & Monitoring Failures','✓ Full audit logging implemented'),
            ('A10 — Server-Side Request Forgery (SSRF)',    '✓ No outbound server-initiated requests to user-supplied URLs'),
        ],
        col_widths=[7.5, 9.0]
    )
    doc.add_paragraph()

    # ── 11. DATA ISOLATION ────────────────────────────────────────────────────
    section_heading(doc, '11. Student Data Isolation & Privacy')
    bullet_item(doc, 'No student accounts:', 'students are enrolled by ID number only — no passwords, no student logins, no student-facing interface')
    bullet_item(doc, 'Minimum PII collected:', 'student ID, name, year of study only — no medical records, contact details, or sensitive information')
    bullet_item(doc, 'Examiner–student separation:', 'an examiner sees only the student at their assigned station and their own scoring form — no visibility into other stations')
    bullet_item(doc, 'Cross-department blocking:', 'coordinators from Department A are physically blocked (at both app and DB levels) from viewing Department B\'s student data')
    bullet_item(doc, 'No public indexing:', 'X-Robots-Tag: noindex on every page — exam URLs cannot be crawled by search engines')
    bullet_item(doc, 'No unauthenticated endpoints:', 'every URL in the system requires a valid session — there are zero public data routes')
    bullet_item(doc, 'Search engine blocking middleware:', 'SearchEngineBlockingMiddleware adds X-Robots-Tag header to all responses automatically')
    doc.add_paragraph()

    # ── 12. QUICK REFERENCE ───────────────────────────────────────────────────
    section_heading(doc, '12. Security Quick-Reference — Key Facts')
    styled_table(doc,
        ['Security Feature', 'Implementation Detail'],
        [
            ('Authentication method',       'Session-based (HttpOnly, SameSite=Lax, Secure in prod)'),
            ('CSRF protection',             'Django CSRF tokens on all state-changing requests'),
            ('Brute-force protection',      'django-axes 8.3 — lockout by IP + username'),
            ('Password minimum',            '8 characters; rejects common, numeric-only, username-similar'),
            ('First-login force-change',    'Middleware-enforced — blocks all pages and APIs until changed'),
            ('Session timeout (admin)',     '10 minutes inactivity auto-logout'),
            ('Session timeout (examiner)', '20 minutes inactivity auto-logout'),
            ('SQL injection protection',    '100% Django ORM parameterized queries — zero raw SQL'),
            ('Access control layers',       'Middleware → DRF permission classes → scope_queryset() → PostgreSQL RLS'),
            ('Database row-level security', 'PostgreSQL RLS policies on all core tables — engine-level enforcement'),
            ('Admin panel security',        'Obscured URL + session gate (double-lock)'),
            ('Transport security',          'HTTPS enforced, HSTS 1 year, SSL DB connections'),
            ('HTTP security headers',       'CSP, X-Frame-Options, nosniff, Referrer-Policy, Permissions-Policy'),
            ('Audit logging',               'Login attempts, unauthorized access, scores, session lifecycle'),
            ('Logging mechanism',           'Celery async queue — does not slow exam operations'),
            ('Soft-delete protection',      'Deactivated accounts explicitly blocked at login'),
            ('Student privacy',             'ID number only; no student logins; department-isolated'),
            ('Last security audit',         'March 13, 2026 — all critical/high findings resolved'),
        ],
        col_widths=[6.0, 10.5]
    )
    doc.add_paragraph()

    # Footer note
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(12)
    nr = note.add_run(
        'This document reflects the actual production implementation. '
        'The full security audit report and source code are available for review upon request.'
    )
    nr.italic = True
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = MID_GREY

    return doc


# ──────────────────────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    output_path = os.path.join(os.path.dirname(__file__), 'OSCE_Security_Brief.docx')
    doc = build_doc()
    doc.save(output_path)
    print(f'Word document generated: {output_path}')
